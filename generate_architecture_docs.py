"""Generate Spanish and English Word architecture documents for CSM-WA."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = Path(__file__).parent / "docs"


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title(doc: Document, text: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(12)
        r2.italic = True
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_diagram_box(doc: Document, title: str, lines: list[str]) -> None:
    add_para(doc, title, bold=True)
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_flow_table(doc: Document, title: str, steps: list[tuple[str, str, str]]) -> None:
    add_para(doc, title, bold=True)
    add_table(doc, ["Origen", "Destino", "Acción"], steps)


def build_document(lang: str) -> Document:
    doc = Document()
    set_doc_defaults(doc)

    if lang == "es":
        add_title(
            doc,
            "CSM-WA — Plan de Arquitectura Híbrida con IA",
            "Versión 1.0 | Julio 2026 | Estado: Planificación",
        )

        add_heading(doc, "1. Resumen ejecutivo")
        add_para(
            doc,
            "CSM-WA es una herramienta web local para equipos de soporte técnico. Hoy busca "
            "workarounds en archivos .txt usando reglas de texto (workaround_parser.py). "
            "La arquitectura propuesta añade una capa de IA que analiza los mejores candidatos, "
            "explica el match y sugiere soluciones — sin reemplazar el parser local.",
        )
        add_diagram_box(
            doc,
            "Modelo híbrido propuesto:",
            [
                "Capa 1 (local, rápida, sin costo)  →  encuentra candidatos",
                "Capa 2 (IA, inteligente, con costo) →  elige, explica y sugiere",
            ],
        )

        add_heading(doc, "2. Arquitectura actual")
        add_para(doc, "Diagrama de componentes y flujo:", bold=True)
        add_diagram_box(
            doc,
            "ESQUEMA — ARQUITECTURA ACTUAL",
            [
                "┌─────────────────────────────────────────────────────────────┐",
                "│                         USUARIO                              │",
                "│              Navegador (http://127.0.0.1:5000)               │",
                "└──────────────────────────┬──────────────────────────────────┘",
                "                           │",
                "                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│                    FRONTEND (local)                          │",
                "│   home.html  |  wa.html  |  procesos.html                   │",
                "│   wa.js / procesos.js / common.js / style.css                │",
                "└──────────────────────────┬──────────────────────────────────┘",
                "                           │ POST /api/wa/search",
                "                           │ POST /api/wa/add",
                "                           │ POST /api/procesos/search",
                "                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│                 BACKEND FLASK (local)                        │",
                "│   app.py                                                     │",
                "│   workaround_parser.py  |  process_parser.py                 │",
                "└──────────────┬─────────────────────────┬────────────────────┘",
                "               │ lee + parsea + busca    │",
                "               ▼                         ▼",
                "┌──────────────────────────┐  ┌──────────────────────────────┐",
                "│ data/workarounds/*.txt   │  │ data/procesos/*.txt          │",
                "│ ~127 Work Arounds        │  │ ~33 Procesos                 │",
                "└──────────────────────────┘  └──────────────────────────────┘",
                "               │",
                "               ▼ JSON resultados",
                "┌─────────────────────────────────────────────────────────────┐",
                "│   Frontend muestra tarjetas acordeón con resultados          │",
                "└─────────────────────────────────────────────────────────────┘",
            ],
        )

        add_heading(doc, "2.1 Flujo actual", level=2)
        add_numbered(
            doc,
            [
                "El usuario escribe un error o pega un Tuxedo completo.",
                "wa.js envía la consulta a POST /api/wa/search.",
                "Flask carga todos los archivos .txt de data/workarounds/.",
                "workaround_parser.py parsea entradas y aplica reglas de coincidencia.",
                "Devuelve lista ordenada por score.",
                "El frontend muestra resultados; el usuario expande manualmente.",
            ],
        )

        add_heading(doc, "2.2 Componentes actuales", level=2)
        add_table(
            doc,
            ["Componente", "Rol"],
            [
                ["app.py", "Servidor Flask, rutas web y API REST"],
                ["workaround_parser.py", "Parser, búsqueda y formato de WAs"],
                ["process_parser.py", "Parser, búsqueda y formato de procesos"],
                ["templates/", "Interfaz HTML"],
                ["static/js/", "Búsqueda, formularios, acordeones"],
                ["data/workarounds/", "Archivos .txt con workarounds"],
                ["data/procesos/", "Archivos .txt con procedimientos"],
            ],
        )

        add_heading(doc, "2.3 Limitaciones actuales", level=2)
        add_bullets(
            doc,
            [
                "No explica por qué un WA es el mejor match.",
                "No combina información de varios WAs.",
                "No sugiere soluciones nuevas si ningún WA encaja.",
                "Búsqueda basada en reglas, no razonamiento semántico profundo.",
            ],
        )

        add_heading(doc, "3. Arquitectura nueva (objetivo)")
        add_para(doc, "Diagrama de la arquitectura híbrida propuesta:", bold=True)
        add_diagram_box(
            doc,
            "ESQUEMA — ARQUITECTURA HÍBRIDA CON IA",
            [
                "┌─────────────────────────────────────────────────────────────┐",
                "│                         USUARIO                              │",
                "│         Navegador + Toggle: Local / Híbrido / IA             │",
                "└──────────────────────────┬──────────────────────────────────┘",
                "                           │",
                "                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│              FRONTEND (wa.html + panel IA)                   │",
                "│   Búsqueda | Recomendación IA | Resultados locales           │",
                "└──────────────────────────┬──────────────────────────────────┘",
                "                           │ POST /api/wa/ai-search",
                "                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│                    BACKEND FLASK                             │",
                "│                                                              │",
                "│  ┌─────────────────────┐    ┌──────────────────────────┐  │",
                "│  │ workaround_parser.py│───▶│ ai_search_service.py     │  │",
                "│  │ (pre-filtro local)  │    │ (orquestador)            │  │",
                "│  │ top 5-10 candidatos │    │                          │  │",
                "│  └─────────────────────┘    │  context_builder.py    │  │",
                "│                              │  prompt_builder.py       │  │",
                "│                              │  search_cache.py         │  │",
                "│                              └───────────┬──────────────┘  │",
                "└──────────────────────────────────────────┼─────────────────┘",
                "                                           │ prompt + candidatos",
                "                                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│              SERVICIO IA (externo)                           │",
                "│   Cursor SDK / OpenAI / Azure OpenAI                         │",
                "│   → ranking + explicación + sugerencia (JSON)                │",
                "└──────────────────────────┬──────────────────────────────────┘",
                "                           │",
                "                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│  Respuesta: WA recomendado + alternativas + explicación      │",
                "│  Fallback automático al parser local si IA falla             │",
                "└─────────────────────────────────────────────────────────────┘",
            ],
        )

        add_heading(doc, "3.1 Flujo nuevo", level=2)
        add_numbered(
            doc,
            [
                "Usuario pega error o Tuxedo.",
                "Fase 1 (local): workaround_parser.py hace pre-filtro → top 5-10 candidatos.",
                "Fase 2 (IA): se construye prompt con query + candidatos (no los 127 WAs).",
                "IA devuelve: WA recomendado, por qué encaja, alternativas, sugerencia.",
                "Frontend muestra recomendación IA + resultados locales como respaldo.",
            ],
        )

        add_heading(doc, "3.2 Propuesta de interfaz", level=2)
        add_diagram_box(
            doc,
            "Vista del buscador en modo híbrido:",
            [
                "[ Caja de búsqueda ]",
                "[ Modo: ● Híbrido   ○ Solo local ]",
                "",
                "┌─ Recomendación IA ─────────────────────────┐",
                "│ Mejor match: RESUME                          │",
                "│ Confianza: Alta                              │",
                "│ Por qué: Coincide csRsCanSub, SERVICE_...    │",
                "│ [Ver WA completo]                            │",
                "└──────────────────────────────────────────────┘",
                "",
                "Otras coincidencias posibles (búsqueda local)",
                "  ▼ RESUME SUBSCRIBER",
                "  ▼ CANCELATION",
            ],
        )

        add_heading(doc, "4. Comparación")
        add_table(
            doc,
            ["Aspecto", "Actual", "Híbrida"],
            [
                ["Motor", "Solo parser local", "Parser + IA"],
                ["Velocidad", "Milisegundos", "Segundos (con IA)"],
                ["Costo", "$0", "Tokens / API"],
                ["Internet", "No requerida", "Requerida para IA"],
                ["Explicación", "No", "Sí"],
                ["Sugerencia nueva", "No", "Sí"],
                ["Privacidad", "100% local", "Datos a servicio externo"],
                ["Fallback", "N/A", "Automático al parser"],
            ],
        )

        add_heading(doc, "5. Plan de implementación")
        phases_es = [
            ("Fase 0 — Decisiones (1-2 días)", "Elegir proveedor IA, política de datos, presupuesto, contrato JSON."),
            ("Fase 1 — Backend IA (3-5 días)", "ai_search_service.py, context_builder.py, prompt_builder.py, /api/wa/ai-search."),
            ("Fase 2 — Pre-filtro (1-2 días)", "Afinar parser, top_n configurable, metadata de candidatos."),
            ("Fase 3 — Frontend (2-3 días)", "Toggle modo, panel IA, explicación, loading, fallback visual."),
            ("Fase 4 — Seguridad (1-2 días)", "Enmascarar BAN/CTN, .env, rate limiting, auditoría."),
            ("Fase 5 — Procesos (2-3 días)", "Extender arquitectura a /procesos (opcional)."),
            ("Fase 6 — Optimizaciones (1-2 sem)", "Embeddings, cache, feedback, guardar sugerencias."),
        ]
        for title, detail in phases_es:
            add_heading(doc, title, level=2)
            add_para(doc, detail)

        add_heading(doc, "6. Contrato API propuesto")
        add_para(doc, "Endpoint: POST /api/wa/ai-search", bold=True)
        add_para(doc, "Request:")
        add_diagram_box(
            doc,
            "",
            [
                '{',
                '  "query": "1) Failed to retrieve array from SERVICE_AGREEMENT...",',
                '  "mode": "hybrid",',
                '  "top_n": 8',
                '}',
            ],
        )
        add_para(doc, "Response (campos principales):")
        add_table(
            doc,
            ["Campo", "Descripción"],
            [
                ["recommended", "WA más relevante según IA"],
                ["confidence", "high | medium | low"],
                ["reason", "Explicación del match"],
                ["alternatives", "Otros WAs posibles"],
                ["suggested_solution", "Solución nueva si no hay match perfecto"],
                ["fallback_used", "Si se usó parser local por fallo de IA"],
            ],
        )

        add_heading(doc, "7. Estructura de archivos nueva")
        add_bullets(
            doc,
            [
                "ai_search_service.py — orquestador de llamadas IA",
                "context_builder.py — arma contexto con candidatos",
                "prompt_builder.py — plantillas de prompt",
                "ai_response_parser.py — valida JSON de IA",
                "config/ai_settings.json — configuración",
                ".env — API keys (fuera del repo)",
            ],
        )

        add_heading(doc, "8. Estimación de esfuerzo")
        add_para(doc, "MVP híbrido funcional: aproximadamente 2 semanas.")

        add_heading(doc, "9. Riesgos y mitigaciones")
        add_table(
            doc,
            ["Riesgo", "Mitigación"],
            [
                ["IA inventa pasos SQL", "Prompt estricto: solo usar WAs proporcionados"],
                ["Costo alto", "Pre-filtro local + cache + límite diario"],
                ["Latencia", "Mostrar resultados locales primero"],
                ["Datos sensibles", "Enmascarar BAN/CTN antes de enviar"],
                ["IA no disponible", "Fallback automático al parser"],
            ],
        )

        add_heading(doc, "10. Recomendación final")
        add_bullets(
            doc,
            [
                "No reemplazar workaround_parser.py.",
                "Usarlo como capa 1 (pre-filtro rápido y gratuito).",
                "Añadir IA como capa 2 (razonamiento y explicación).",
                "Siempre tener fallback al parser local.",
                "Enmascarar datos sensibles antes de llamadas externas.",
                "Validar que la IA solo recomiende WAs existentes en archivos locales.",
            ],
        )

        add_flow_table(
            doc,
            "Flujo de datos — modo híbrido",
            [
                ("Usuario", "wa.js", "Pega error Tuxedo"),
                ("wa.js", "app.py", "POST /api/wa/ai-search"),
                ("app.py", "workaround_parser.py", "Pre-filtro local"),
                ("workaround_parser.py", "ai_search_service.py", "Top N candidatos"),
                ("ai_search_service.py", "Servicio IA", "Prompt + contexto"),
                ("Servicio IA", "ai_search_service.py", "JSON recomendación"),
                ("app.py", "wa.js", "Recomendación + resultados"),
                ("wa.js", "Usuario", "Muestra mejor match + alternativas"),
            ],
        )

    else:
        add_title(
            doc,
            "CSM-WA — Hybrid AI Architecture Plan",
            "Version 1.0 | July 2026 | Status: Planning",
        )

        add_heading(doc, "1. Executive summary")
        add_para(
            doc,
            "CSM-WA is a local web tool for technical support teams. Today it searches "
            "workarounds in .txt files using text rules (workaround_parser.py). "
            "The proposed architecture adds an AI layer that analyzes the best candidates, "
            "explains the match, and suggests solutions — without replacing the local parser.",
        )
        add_diagram_box(
            doc,
            "Proposed hybrid model:",
            [
                "Layer 1 (local, fast, no cost)   →  finds candidates",
                "Layer 2 (AI, smart, paid)          →  picks, explains, suggests",
            ],
        )

        add_heading(doc, "2. Current architecture")
        add_para(doc, "Component diagram and flow:", bold=True)
        add_diagram_box(
            doc,
            "DIAGRAM — CURRENT ARCHITECTURE",
            [
                "┌─────────────────────────────────────────────────────────────┐",
                "│                           USER                               │",
                "│               Browser (http://127.0.0.1:5000)                │",
                "└──────────────────────────┬──────────────────────────────────┘",
                "                           │",
                "                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│                     FRONTEND (local)                         │",
                "│   home.html  |  wa.html  |  procesos.html                    │",
                "│   wa.js / procesos.js / common.js / style.css                │",
                "└──────────────────────────┬──────────────────────────────────┘",
                "                           │ POST /api/wa/search",
                "                           │ POST /api/wa/add",
                "                           │ POST /api/procesos/search",
                "                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│                  FLASK BACKEND (local)                       │",
                "│   app.py                                                     │",
                "│   workaround_parser.py  |  process_parser.py                 │",
                "└──────────────┬─────────────────────────┬────────────────────┘",
                "               │ read + parse + search   │",
                "               ▼                         ▼",
                "┌──────────────────────────┐  ┌──────────────────────────────┐",
                "│ data/workarounds/*.txt   │  │ data/procesos/*.txt          │",
                "│ ~127 Work Arounds        │  │ ~33 Processes                │",
                "└──────────────────────────┘  └──────────────────────────────┘",
                "               │",
                "               ▼ JSON results",
                "┌─────────────────────────────────────────────────────────────┐",
                "│   Frontend displays accordion cards with results             │",
                "└─────────────────────────────────────────────────────────────┘",
            ],
        )

        add_heading(doc, "2.1 Current flow", level=2)
        add_numbered(
            doc,
            [
                "User types an error or pastes a full Tuxedo stack trace.",
                "wa.js sends the query to POST /api/wa/search.",
                "Flask loads all .txt files from data/workarounds/.",
                "workaround_parser.py parses entries and applies matching rules.",
                "Returns a list sorted by score.",
                "Frontend displays results; user expands manually.",
            ],
        )

        add_heading(doc, "2.2 Current components", level=2)
        add_table(
            doc,
            ["Component", "Role"],
            [
                ["app.py", "Flask server, web routes and REST API"],
                ["workaround_parser.py", "Parser, search, and WA formatting"],
                ["process_parser.py", "Parser, search, and process formatting"],
                ["templates/", "HTML UI"],
                ["static/js/", "Search, forms, accordions"],
                ["data/workarounds/", ".txt files with workarounds"],
                ["data/procesos/", ".txt files with procedures"],
            ],
        )

        add_heading(doc, "2.3 Current limitations", level=2)
        add_bullets(
            doc,
            [
                "Does not explain why a WA is the best match.",
                "Cannot combine information from multiple WAs.",
                "Does not suggest new solutions when no WA fits.",
                "Rule-based search, not deep semantic reasoning.",
            ],
        )

        add_heading(doc, "3. Target architecture")
        add_para(doc, "Proposed hybrid architecture diagram:", bold=True)
        add_diagram_box(
            doc,
            "DIAGRAM — HYBRID AI ARCHITECTURE",
            [
                "┌─────────────────────────────────────────────────────────────┐",
                "│                           USER                               │",
                "│          Browser + Toggle: Local / Hybrid / AI               │",
                "└──────────────────────────┬──────────────────────────────────┘",
                "                           │",
                "                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│               FRONTEND (wa.html + AI panel)                  │",
                "│   Search | AI Recommendation | Local results               │",
                "└──────────────────────────┬──────────────────────────────────┘",
                "                           │ POST /api/wa/ai-search",
                "                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│                    FLASK BACKEND                             │",
                "│                                                              │",
                "│  ┌─────────────────────┐    ┌──────────────────────────┐  │",
                "│  │ workaround_parser.py│───▶│ ai_search_service.py     │  │",
                "│  │ (local pre-filter)  │    │ (orchestrator)           │  │",
                "│  │ top 5-10 candidates │    │                          │  │",
                "│  └─────────────────────┘    │  context_builder.py    │  │",
                "│                              │  prompt_builder.py       │  │",
                "│                              │  search_cache.py         │  │",
                "│                              └───────────┬──────────────┘  │",
                "└──────────────────────────────────────────┼─────────────────┘",
                "                                           │ prompt + candidates",
                "                                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│                AI SERVICE (external)                         │",
                "│   Cursor SDK / OpenAI / Azure OpenAI                         │",
                "│   → ranking + explanation + suggestion (JSON)                │",
                "└──────────────────────────┬──────────────────────────────────┘",
                "                           │",
                "                           ▼",
                "┌─────────────────────────────────────────────────────────────┐",
                "│  Response: recommended WA + alternatives + explanation       │",
                "│  Automatic fallback to local parser if AI fails              │",
                "└─────────────────────────────────────────────────────────────┘",
            ],
        )

        add_heading(doc, "3.1 New flow", level=2)
        add_numbered(
            doc,
            [
                "User pastes error or Tuxedo stack trace.",
                "Phase 1 (local): workaround_parser.py pre-filters → top 5-10 candidates.",
                "Phase 2 (AI): prompt built with query + candidates (not all 127 WAs).",
                "AI returns: recommended WA, why it fits, alternatives, suggestion.",
                "Frontend shows AI recommendation + local results as backup.",
            ],
        )

        add_heading(doc, "3.2 Proposed UI", level=2)
        add_diagram_box(
            doc,
            "Hybrid search view:",
            [
                "[ Search box ]",
                "[ Mode: ● Hybrid   ○ Local only ]",
                "",
                "┌─ AI Recommendation ──────────────────────────┐",
                "│ Best match: RESUME                             │",
                "│ Confidence: High                               │",
                "│ Why: Matches csRsCanSub, SERVICE_...           │",
                "│ [View full WA]                                 │",
                "└────────────────────────────────────────────────┘",
                "",
                "Other possible matches (local search)",
                "  ▼ RESUME SUBSCRIBER",
                "  ▼ CANCELATION",
            ],
        )

        add_heading(doc, "4. Comparison")
        add_table(
            doc,
            ["Aspect", "Current", "Hybrid"],
            [
                ["Engine", "Local parser only", "Parser + AI"],
                ["Speed", "Milliseconds", "Seconds (with AI)"],
                ["Cost", "$0", "Tokens / API"],
                ["Internet", "Not required", "Required for AI"],
                ["Explanation", "No", "Yes"],
                ["New suggestion", "No", "Yes"],
                ["Privacy", "Fully local", "Data to external service"],
                ["Fallback", "N/A", "Automatic to parser"],
            ],
        )

        add_heading(doc, "5. Implementation plan")
        phases_en = [
            ("Phase 0 — Decisions (1-2 days)", "Choose AI provider, data policy, budget, JSON contract."),
            ("Phase 1 — AI backend (3-5 days)", "ai_search_service.py, context_builder.py, prompt_builder.py, /api/wa/ai-search."),
            ("Phase 2 — Pre-filter (1-2 days)", "Tune parser, configurable top_n, candidate metadata."),
            ("Phase 3 — Frontend (2-3 days)", "Mode toggle, AI panel, explanation, loading, visual fallback."),
            ("Phase 4 — Security (1-2 days)", "Mask BAN/CTN, .env, rate limiting, audit."),
            ("Phase 5 — Processes (2-3 days)", "Extend architecture to /procesos (optional)."),
            ("Phase 6 — Optimizations (1-2 wks)", "Embeddings, cache, feedback, save suggestions."),
        ]
        for title, detail in phases_en:
            add_heading(doc, title, level=2)
            add_para(doc, detail)

        add_heading(doc, "6. Proposed API contract")
        add_para(doc, "Endpoint: POST /api/wa/ai-search", bold=True)
        add_para(doc, "Request:")
        add_diagram_box(
            doc,
            "",
            [
                '{',
                '  "query": "1) Failed to retrieve array from SERVICE_AGREEMENT...",',
                '  "mode": "hybrid",',
                '  "top_n": 8',
                '}',
            ],
        )
        add_para(doc, "Response (main fields):")
        add_table(
            doc,
            ["Field", "Description"],
            [
                ["recommended", "Most relevant WA per AI"],
                ["confidence", "high | medium | low"],
                ["reason", "Match explanation"],
                ["alternatives", "Other possible WAs"],
                ["suggested_solution", "New solution if no perfect match"],
                ["fallback_used", "Whether local parser was used due to AI failure"],
            ],
        )

        add_heading(doc, "7. New file structure")
        add_bullets(
            doc,
            [
                "ai_search_service.py — AI call orchestrator",
                "context_builder.py — builds context with candidates",
                "prompt_builder.py — prompt templates",
                "ai_response_parser.py — validates AI JSON",
                "config/ai_settings.json — configuration",
                ".env — API keys (outside repo)",
            ],
        )

        add_heading(doc, "8. Effort estimate")
        add_para(doc, "Functional hybrid MVP: approximately 2 weeks.")

        add_heading(doc, "9. Risks and mitigations")
        add_table(
            doc,
            ["Risk", "Mitigation"],
            [
                ["AI invents SQL steps", "Strict prompt: only use provided WAs"],
                ["High cost", "Local pre-filter + cache + daily limit"],
                ["Latency", "Show local results first"],
                ["Sensitive data", "Mask BAN/CTN before sending"],
                ["AI unavailable", "Automatic fallback to parser"],
            ],
        )

        add_heading(doc, "10. Final recommendation")
        add_bullets(
            doc,
            [
                "Do not replace workaround_parser.py.",
                "Use it as layer 1 (fast, free pre-filter).",
                "Add AI as layer 2 (reasoning and explanation).",
                "Always have fallback to local parser.",
                "Mask sensitive data before external calls.",
                "Validate AI only recommends WAs existing in local files.",
            ],
        )

        add_table(
            doc,
            ["From", "To", "Action"],
            [
                ("User", "wa.js", "Paste Tuxedo error"),
                ("wa.js", "app.py", "POST /api/wa/ai-search"),
                ("app.py", "workaround_parser.py", "Local pre-filter"),
                ("workaround_parser.py", "ai_search_service.py", "Top N candidates"),
                ("ai_search_service.py", "AI service", "Prompt + context"),
                ("AI service", "ai_search_service.py", "JSON recommendation"),
                ("app.py", "wa.js", "Recommendation + results"),
                ("wa.js", "User", "Show best match + alternatives"),
            ],
        )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "CSM-WA — Planning document. No code changes until approved."
        if lang == "en"
        else "CSM-WA — Documento de planificación. Sin cambios en código hasta aprobación."
    )
    run.italic = True
    run.font.size = Pt(9)

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    es_path = OUTPUT_DIR / "CSM-WA_Arquitectura_IA_Hibrida_ES.docx"
    en_path = OUTPUT_DIR / "CSM-WA_Hybrid_AI_Architecture_EN.docx"

    build_document("es").save(es_path)
    build_document("en").save(en_path)

    print(f"Created: {es_path}")
    print(f"Created: {en_path}")


if __name__ == "__main__":
    main()
